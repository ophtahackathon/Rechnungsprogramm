import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import sqlite3
from datetime import datetime
from PIL import Image, ImageTk
import os
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from tkcalendar import DateEntry

class Rechnungsprogramm:
    def __init__(self, root):
        self.root = root
        self.root.title("Rechnungsprogramm")
        self.root.geometry("900x700")
        self.root.configure(bg='#2c2c2c')  # Dunkler Hintergrund

        # Stil für dunkles Design
        style = ttk.Style()
        style.configure('TNotebook', background='#2c2c2c', borderwidth=0)
        style.configure('TNotebook.Tab', background='#404040', foreground='white', padding=[10, 5])
        style.map('TNotebook.Tab', background=[('selected', '#555555')])
        style.configure('TLabel', background='#2c2c2c', foreground='white')
        style.configure('TEntry', fieldbackground='#404040', foreground='white')
        style.configure('TButton', background='#555555', foreground='white')
        style.configure('Treeview', background='#404040', foreground='white', fieldbackground='#404040')
        style.configure('Treeview.Heading', background='#555555', foreground='white')

        # Datenbank verbinden
        self.conn = sqlite3.connect('rechnungen.db')
        self.c = self.conn.cursor()

        # Tabs erstellen
        self.tab_control = ttk.Notebook(root)
        self.tab_firma = ttk.Frame(self.tab_control, style='TFrame')
        self.tab_mandanten = ttk.Frame(self.tab_control, style='TFrame')
        self.tab_leistungen = ttk.Frame(self.tab_control, style='TFrame')
        self.tab_termine = ttk.Frame(self.tab_control, style='TFrame')
        self.tab_rechnungen = ttk.Frame(self.tab_control, style='TFrame')
        self.tab_statistik = ttk.Frame(self.tab_control, style='TFrame')

        self.tab_control.add(self.tab_firma, text='Firma')
        self.tab_control.add(self.tab_mandanten, text='Mandanten')
        self.tab_control.add(self.tab_leistungen, text='Leistungen')
        self.tab_control.add(self.tab_termine, text='Termine')
        self.tab_control.add(self.tab_rechnungen, text='Rechnungen')
        self.tab_control.add(self.tab_statistik, text='Statistik')

        self.tab_control.pack(expand=1, fill="both")

        # Tabs initialisieren
        self.init_firma_tab()
        self.init_mandanten_tab()
        self.init_leistungen_tab()
        self.init_termine_tab()
        self.init_rechnungen_tab()
        self.init_statistik_tab()

    def init_firma_tab(self):
        # Firma Daten laden
        self.c.execute("SELECT * FROM firma WHERE id=1")
        firma = self.c.fetchone()
        if not firma:
            self.c.execute("INSERT INTO firma (id, name, adresse, logo_path, mwst_satz, mwst_text) VALUES (1, '', '', '', 19.0, 'Es wird keine Mehrwertsteuer erhoben')")
            self.conn.commit()
            firma = (1, '', '', '', 19.0, 'Es wird keine Mehrwertsteuer erhoben')
        elif firma[1] is None or firma[2] is None or firma[3] is None or firma[5] is None:
            # Falls Felder NULL sind, mit Defaults füllen
            self.c.execute("UPDATE firma SET name=COALESCE(name, ''), adresse=COALESCE(adresse, ''), logo_path=COALESCE(logo_path, ''), mwst_text=COALESCE(mwst_text, 'Es wird keine Mehrwertsteuer erhoben') WHERE id=1")
            self.conn.commit()
            self.c.execute("SELECT * FROM firma WHERE id=1")
            firma = self.c.fetchone()

        # Widgets
        ttk.Label(self.tab_firma, text="Firmenname:").grid(row=0, column=0, sticky=tk.W)
        self.firma_name = ttk.Entry(self.tab_firma)
        self.firma_name.grid(row=0, column=1)
        self.firma_name.insert(0, firma[1] or '')

        ttk.Label(self.tab_firma, text="Adresse:").grid(row=1, column=0, sticky=tk.W)
        self.firma_adresse = tk.Text(self.tab_firma, height=3, width=30)
        self.firma_adresse.grid(row=1, column=1)
        self.firma_adresse.insert(tk.END, firma[2] or '')

        ttk.Label(self.tab_firma, text="Logo:").grid(row=2, column=0, sticky=tk.W)
        self.logo_path = tk.StringVar(value=firma[3] or '')
        ttk.Entry(self.tab_firma, textvariable=self.logo_path).grid(row=2, column=1)
        ttk.Button(self.tab_firma, text="Auswählen", command=self.select_logo).grid(row=2, column=2)

        ttk.Label(self.tab_firma, text="MwSt Satz (%):").grid(row=3, column=0, sticky=tk.W)
        self.mwst_satz = ttk.Entry(self.tab_firma)
        self.mwst_satz.grid(row=3, column=1)
        self.mwst_satz.insert(0, str(firma[4] or 19.0))

        ttk.Label(self.tab_firma, text="MwSt Text:").grid(row=4, column=0, sticky=tk.W)
        self.mwst_text = ttk.Entry(self.tab_firma)
        self.mwst_text.grid(row=4, column=1)
        self.mwst_text.insert(0, firma[5] or 'Es wird keine Mehrwertsteuer erhoben')

    def select_logo(self):
        filename = filedialog.askopenfilename(title="Logo auswählen", filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif")])
        if filename:
            self.logo_path.set(filename)

    def select_vorlage(self):
        filename = filedialog.askopenfilename(title="Vorlage auswählen", filetypes=[("Word files", "*.docx")])
        if filename:
            self.vorlage_path.set(filename)

    def save_firma(self):
        name = self.firma_name.get()
        adresse = self.firma_adresse.get("1.0", tk.END).strip()
        logo = self.logo_path.get()
        mwst = float(self.mwst_satz.get())
        text = self.mwst_text.get()
        self.c.execute("UPDATE firma SET name=?, adresse=?, logo_path=?, mwst_satz=?, mwst_text=? WHERE id=1",
                       (name, adresse, logo, mwst, text))
        self.conn.commit()
        messagebox.showinfo("Erfolg", "Firmendaten gespeichert")

    def save_firma(self):
        name = self.firma_name.get()
        adresse = self.firma_adresse.get("1.0", tk.END).strip()
        logo = self.logo_path.get()
        mwst = float(self.mwst_satz.get())
        text = self.mwst_text.get()
        self.c.execute("UPDATE firma SET name=?, adresse=?, logo_path=?, mwst_satz=?, mwst_text=? WHERE id=1",
                       (name, adresse, logo, mwst, text))
        self.conn.commit()
        messagebox.showinfo("Erfolg", "Firmendaten gespeichert")

    def init_mandanten_tab(self):
        # Liste der Mandanten
        self.mandanten_tree = ttk.Treeview(self.tab_mandanten, columns=("ID", "Name", "Adresse"), show="headings")
        self.mandanten_tree.heading("ID", text="ID")
        self.mandanten_tree.heading("Name", text="Name")
        self.mandanten_tree.heading("Adresse", text="Adresse")
        self.mandanten_tree.pack(fill=tk.BOTH, expand=True)

        # Buttons
        btn_frame = ttk.Frame(self.tab_mandanten)
        btn_frame.pack(fill=tk.X)
        ttk.Button(btn_frame, text="Hinzufügen", command=self.add_mandant).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Bearbeiten", command=self.edit_mandant).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Löschen", command=self.delete_mandant).pack(side=tk.LEFT)

        self.load_mandanten()

    def load_mandanten(self):
        for i in self.mandanten_tree.get_children():
            self.mandanten_tree.delete(i)
        self.c.execute("SELECT id, name, adresse FROM mandanten")
        for row in self.c.fetchall():
            self.mandanten_tree.insert("", tk.END, values=row)

    def add_mandant(self):
        self.mandant_dialog("Hinzufügen")

    def edit_mandant(self):
        selected = self.mandanten_tree.selection()
        if not selected:
            messagebox.showerror("Fehler", "Bitte Mandant auswählen")
            return
        item = self.mandanten_tree.item(selected[0])
        id_ = item['values'][0]
        self.c.execute("SELECT * FROM mandanten WHERE id=?", (id_,))
        mandant = self.c.fetchone()
        self.mandant_dialog("Bearbeiten", mandant)

    def delete_mandant(self):
        selected = self.mandanten_tree.selection()
        if not selected:
            messagebox.showerror("Fehler", "Bitte Mandant auswählen")
            return
        item = self.mandanten_tree.item(selected[0])
        id_ = item['values'][0]
        self.c.execute("DELETE FROM mandanten WHERE id=?", (id_,))
        self.conn.commit()
        self.load_mandanten()

    def mandant_dialog(self, action, mandant=None):
        dialog = tk.Toplevel(self.root)
        dialog.title(f"Mandant {action}")
        dialog.geometry("400x300")

        ttk.Label(dialog, text="Name:").grid(row=0, column=0)
        name_entry = ttk.Entry(dialog)
        name_entry.grid(row=0, column=1)

        ttk.Label(dialog, text="Adresse:").grid(row=1, column=0)
        adresse_entry = tk.Text(dialog, height=3, width=30)
        adresse_entry.grid(row=1, column=1)

        ttk.Label(dialog, text="Email:").grid(row=2, column=0)
        email_entry = ttk.Entry(dialog)
        email_entry.grid(row=2, column=1)

        ttk.Label(dialog, text="Telefon:").grid(row=3, column=0)
        telefon_entry = ttk.Entry(dialog)
        telefon_entry.grid(row=3, column=1)

        if mandant:
            name_entry.insert(0, mandant[1])
            adresse_entry.insert(tk.END, mandant[2])
            email_entry.insert(0, mandant[3])
            telefon_entry.insert(0, mandant[4])

        def save():
            name = name_entry.get()
            adresse = adresse_entry.get("1.0", tk.END).strip()
            email = email_entry.get()
            telefon = telefon_entry.get()
            if action == "Hinzufügen":
                self.c.execute("INSERT INTO mandanten (name, adresse, email, telefon) VALUES (?, ?, ?, ?)",
                               (name, adresse, email, telefon))
            else:
                self.c.execute("UPDATE mandanten SET name=?, adresse=?, email=?, telefon=? WHERE id=?",
                               (name, adresse, email, telefon, mandant[0]))
            self.conn.commit()
            self.load_mandanten()
            dialog.destroy()

        ttk.Button(dialog, text="Speichern", command=save).grid(row=4, column=0, columnspan=2)

    # Ähnlich für Leistungen
    def init_leistungen_tab(self):
        self.leistungen_tree = ttk.Treeview(self.tab_leistungen, columns=("ID", "Kürzel", "Beschreibung", "Betrag"), show="headings")
        self.leistungen_tree.heading("ID", text="ID")
        self.leistungen_tree.heading("Kürzel", text="Kürzel")
        self.leistungen_tree.heading("Beschreibung", text="Beschreibung")
        self.leistungen_tree.heading("Betrag", text="Betrag")
        self.leistungen_tree.pack(fill=tk.BOTH, expand=True)

        btn_frame = ttk.Frame(self.tab_leistungen)
        btn_frame.pack(fill=tk.X)
        ttk.Button(btn_frame, text="Hinzufügen", command=self.add_leistung).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Bearbeiten", command=self.edit_leistung).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Löschen", command=self.delete_leistung).pack(side=tk.LEFT)

        self.load_leistungen()

    def load_leistungen(self):
        for i in self.leistungen_tree.get_children():
            self.leistungen_tree.delete(i)
        self.c.execute("SELECT id, kuerzel, beschreibung, betrag FROM leistungen")
        for row in self.c.fetchall():
            self.leistungen_tree.insert("", tk.END, values=row)

    def add_leistung(self):
        self.leistung_dialog("Hinzufügen")

    def edit_leistung(self):
        selected = self.leistungen_tree.selection()
        if not selected:
            messagebox.showerror("Fehler", "Bitte Leistung auswählen")
            return
        item = self.leistungen_tree.item(selected[0])
        id_ = item['values'][0]
        self.c.execute("SELECT * FROM leistungen WHERE id=?", (id_,))
        leistung = self.c.fetchone()
        self.leistung_dialog("Bearbeiten", leistung)

    def delete_leistung(self):
        selected = self.leistungen_tree.selection()
        if not selected:
            messagebox.showerror("Fehler", "Bitte Leistung auswählen")
            return
        item = self.leistungen_tree.item(selected[0])
        id_ = item['values'][0]
        self.c.execute("DELETE FROM leistungen WHERE id=?", (id_,))
        self.conn.commit()
        self.load_leistungen()

    def leistung_dialog(self, action, leistung=None):
        dialog = tk.Toplevel(self.root)
        dialog.title(f"Leistung {action}")
        dialog.geometry("400x200")

        ttk.Label(dialog, text="Kürzel:").grid(row=0, column=0)
        kuerzel_entry = ttk.Entry(dialog)
        kuerzel_entry.grid(row=0, column=1)

        ttk.Label(dialog, text="Beschreibung:").grid(row=1, column=0)
        beschreibung_entry = ttk.Entry(dialog)
        beschreibung_entry.grid(row=1, column=1)

        ttk.Label(dialog, text="Betrag:").grid(row=2, column=0)
        betrag_entry = ttk.Entry(dialog)
        betrag_entry.grid(row=2, column=1)

        if leistung:
            kuerzel_entry.insert(0, leistung[1])
            beschreibung_entry.insert(0, leistung[2])
            betrag_entry.insert(0, str(leistung[3]))

        def save():
            kuerzel = kuerzel_entry.get()
            beschreibung = beschreibung_entry.get()
            betrag = float(betrag_entry.get())
            if action == "Hinzufügen":
                self.c.execute("INSERT INTO leistungen (kuerzel, beschreibung, betrag) VALUES (?, ?, ?)",
                               (kuerzel, beschreibung, betrag))
            else:
                self.c.execute("UPDATE leistungen SET kuerzel=?, beschreibung=?, betrag=? WHERE id=?",
                               (kuerzel, beschreibung, betrag, leistung[0]))
            self.conn.commit()
            self.load_leistungen()
            dialog.destroy()

        ttk.Button(dialog, text="Speichern", command=save).grid(row=3, column=0, columnspan=2)

    # Für Termine
    def init_termine_tab(self):
        # Mandant auswählen
        ttk.Label(self.tab_termine, text="Mandant:").grid(row=0, column=0)
        self.mandant_var = tk.StringVar()
        self.mandant_combo = ttk.Combobox(self.tab_termine, textvariable=self.mandant_var)
        self.mandant_combo.grid(row=0, column=1)
        self.load_mandanten_combo()

        ttk.Label(self.tab_termine, text="Datum (dd.mm.yyyy):").grid(row=1, column=0, sticky=tk.W)
        self.datum_entry = DateEntry(self.tab_termine, year=datetime.now().year, month=datetime.now().month, 
                                      day=datetime.now().day, dateformat='%d.%m.%Y', locale='de_DE')
        self.datum_entry.grid(row=1, column=1)

        # Leistungen auswählen (DropDown)
        ttk.Label(self.tab_termine, text="Leistung:").grid(row=2, column=0, sticky=tk.W)
        self.leistung_var = tk.StringVar()
        self.leistung_combo = ttk.Combobox(self.tab_termine, textvariable=self.leistung_var)
        self.leistung_combo.grid(row=2, column=1)
        self.load_leistungen_combo()

        ttk.Button(self.tab_termine, text="Termin hinzufügen", command=self.add_termin).grid(row=3, column=0, columnspan=2)

        # Liste der Termine
        self.termine_tree = ttk.Treeview(self.tab_termine, columns=("ID", "Mandant", "Datum", "Leistungen"), show="headings")
        self.termine_tree.heading("ID", text="ID")
        self.termine_tree.heading("Mandant", text="Mandant")
        self.termine_tree.heading("Datum", text="Datum")
        self.termine_tree.heading("Leistungen", text="Leistungen")
        self.termine_tree.grid(row=4, column=0, columnspan=2, sticky="nsew")
        self.tab_termine.grid_rowconfigure(4, weight=1)
        self.tab_termine.grid_columnconfigure(1, weight=1)

        self.load_termine()

    def load_mandanten_combo(self):
        self.c.execute("SELECT id, name FROM mandanten")
        mandanten = self.c.fetchall()
        self.mandant_combo['values'] = [f"{m[0]} - {m[1]}" for m in mandanten]

    def load_leistungen_combo(self):
        self.c.execute("SELECT id, kuerzel FROM leistungen")
        leistungen = self.c.fetchall()
        self.leistung_combo['values'] = [f"{l[0]} - {l[1]}" for l in leistungen]

    def add_termin(self):
        mandant_str = self.mandant_var.get()
        if not mandant_str:
            messagebox.showerror("Fehler", "Mandant auswählen")
            return
        mandant_id = int(mandant_str.split(' - ')[0])
        datum = self.datum_entry.get_date().strftime('%Y-%m-%d')
        leistung_str = self.leistung_var.get()
        if not leistung_str:
            messagebox.showerror("Fehler", "Leistung auswählen")
            return
        leistung_id = int(leistung_str.split(' - ')[0])

        # Termin einfügen
        self.c.execute("INSERT INTO termine (mandant_id, datum) VALUES (?, ?)", (mandant_id, datum))
        termin_id = self.c.lastrowid

        # Leistung zuordnen
        self.c.execute("INSERT INTO termin_leistungen (termin_id, leistung_id) VALUES (?, ?)", (termin_id, leistung_id))

        self.conn.commit()
        self.load_termine()
        messagebox.showinfo("Erfolg", "Termin hinzugefügt")

        self.conn.commit()
        self.load_termine()
        messagebox.showinfo("Erfolg", "Termin hinzugefügt")

    def load_termine(self):
        for i in self.termine_tree.get_children():
            self.termine_tree.delete(i)
        self.c.execute("""
            SELECT t.id, m.name, t.datum, GROUP_CONCAT(l.kuerzel)
            FROM termine t
            JOIN mandanten m ON t.mandant_id = m.id
            LEFT JOIN termin_leistungen tl ON t.id = tl.termin_id
            LEFT JOIN leistungen l ON tl.leistung_id = l.id
            GROUP BY t.id
        """)
        for row in self.c.fetchall():
            self.termine_tree.insert("", tk.END, values=row)

    # Für Rechnungen
    def init_rechnungen_tab(self):
        ttk.Label(self.tab_rechnungen, text="Mandant:").grid(row=0, column=0)
        self.rechnung_mandant_var = tk.StringVar()
        self.rechnung_mandant_combo = ttk.Combobox(self.tab_rechnungen, textvariable=self.rechnung_mandant_var)
        self.rechnung_mandant_combo.grid(row=0, column=1)
        self.load_mandanten_combo_rechnung()

        ttk.Label(self.tab_rechnungen, text="Von Datum (dd.mm.yyyy):").grid(row=1, column=0)
        self.von_datum = DateEntry(self.tab_rechnungen, year=datetime.now().year, month=datetime.now().month,
                                    day=1, dateformat='%d.%m.%Y', locale='de_DE')
        self.von_datum.grid(row=1, column=1)

        ttk.Label(self.tab_rechnungen, text="Bis Datum (dd.mm.yyyy):").grid(row=2, column=0)
        self.bis_datum = DateEntry(self.tab_rechnungen, year=datetime.now().year, month=datetime.now().month,
                                    day=datetime.now().day, dateformat='%d.%m.%Y', locale='de_DE')
        self.bis_datum.grid(row=2, column=1)

        ttk.Button(self.tab_rechnungen, text="Rechnung erstellen", command=self.create_rechnung).grid(row=3, column=0, columnspan=2)

        # Liste der Rechnungen
        self.rechnungen_tree = ttk.Treeview(self.tab_rechnungen, columns=("ID", "Mandant", "Von", "Bis", "Summe"), show="headings")
        self.rechnungen_tree.heading("ID", text="ID")
        self.rechnungen_tree.heading("Mandant", text="Mandant")
        self.rechnungen_tree.heading("Von", text="Von")
        self.rechnungen_tree.heading("Bis", text="Bis")
        self.rechnungen_tree.heading("Summe", text="Summe")
        self.rechnungen_tree.grid(row=4, column=0, columnspan=2, sticky="nsew")
        self.tab_rechnungen.grid_rowconfigure(4, weight=1)
        self.tab_rechnungen.grid_columnconfigure(1, weight=1)

        self.load_rechnungen()

    def load_mandanten_combo_rechnung(self):
        self.c.execute("SELECT id, name FROM mandanten")
        mandanten = self.c.fetchall()
        self.rechnung_mandant_combo['values'] = [f"{m[0]} - {m[1]}" for m in mandanten]

    def create_rechnung(self):
        mandant_str = self.rechnung_mandant_var.get()
        if not mandant_str:
            messagebox.showerror("Fehler", "Mandant auswählen")
            return
        mandant_id = int(mandant_str.split(' - ')[0])
        von = self.von_datum.get_date().strftime('%Y-%m-%d')
        bis = self.bis_datum.get_date().strftime('%Y-%m-%d')

        # Termine im Zeitraum holen
        self.c.execute("""
            SELECT tl.leistung_id, l.beschreibung, l.betrag, COUNT(tl.leistung_id) as anzahl
            FROM termine t
            JOIN termin_leistungen tl ON t.id = tl.termin_id
            JOIN leistungen l ON tl.leistung_id = l.id
            WHERE t.mandant_id = ? AND t.datum BETWEEN ? AND ?
            GROUP BY tl.leistung_id
        """, (mandant_id, von, bis))
        positionen = self.c.fetchall()

        if not positionen:
            messagebox.showerror("Fehler", "Keine Leistungen im Zeitraum")
            return

        summe_netto = sum(p[2] * p[3] for p in positionen)
        self.c.execute("SELECT mwst_satz FROM firma WHERE id=1")
        mwst_satz = self.c.fetchone()[0]
        mwst = summe_netto * (mwst_satz / 100)
        summe_brutto = summe_netto + mwst

        # Rechnung einfügen
        erstellt_am = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.c.execute("INSERT INTO rechnungen (mandant_id, von_datum, bis_datum, summe_netto, mwst, summe_brutto, erstellt_am) VALUES (?, ?, ?, ?, ?, ?, ?)",
                       (mandant_id, von, bis, summe_netto, mwst, summe_brutto, erstellt_am))
        rechnung_id = self.c.lastrowid

        # Positionen einfügen
        for p in positionen:
            self.c.execute("INSERT INTO rechnung_positionen (rechnung_id, leistung_id, anzahl, betrag) VALUES (?, ?, ?, ?)",
                           (rechnung_id, p[0], p[3], p[2]))

        self.conn.commit()
        self.load_rechnungen()
        messagebox.showinfo("Erfolg", f"Rechnung {rechnung_id} erstellt")

        # Export anbieten
        if messagebox.askyesno("Export", "Rechnung als Word-Datei exportieren?"):
            self.export_rechnung(rechnung_id)

    def load_rechnungen(self):
        for i in self.rechnungen_tree.get_children():
            self.rechnungen_tree.delete(i)
        self.c.execute("SELECT r.id, m.name, r.von_datum, r.bis_datum, r.summe_brutto FROM rechnungen r JOIN mandanten m ON r.mandant_id = m.id")
        for row in self.c.fetchall():
            self.rechnungen_tree.insert("", tk.END, values=row)

    def export_rechnung(self, rechnung_id):
        vorlage_path = self.vorlage_path.get()
        if vorlage_path:
            # Vorlage laden
            doc = Document(vorlage_path)
        else:
            # Neue Vorlage erstellen
            doc = Document()
            doc.add_heading('Rechnung', 0)

        # Firma Daten
        self.c.execute("SELECT * FROM firma WHERE id=1")
        firma = self.c.fetchone()

        # Rechnung Daten
        self.c.execute("SELECT * FROM rechnungen WHERE id=?", (rechnung_id,))
        rechnung = self.c.fetchone()
        mandant_id = rechnung[1]
        self.c.execute("SELECT * FROM mandanten WHERE id=?", (mandant_id,))
        mandant = self.c.fetchone()

        # Positionen
        self.c.execute("""
            SELECT l.beschreibung, rp.anzahl, rp.betrag
            FROM rechnung_positionen rp
            JOIN leistungen l ON rp.leistung_id = l.id
            WHERE rp.rechnung_id = ?
        """, (rechnung_id,))
        positionen = self.c.fetchall()

        # Platzhalter ersetzen oder anfügen
        if vorlage_path:
            # Suche nach Platzhaltern und ersetze
            for paragraph in doc.paragraphs:
                if '{{FIRMA_NAME}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{FIRMA_NAME}}', firma[1] or '')
                if '{{FIRMA_ADRESSE}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{FIRMA_ADRESSE}}', firma[2] or '')
                if '{{MANDANT_NAME}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{MANDANT_NAME}}', mandant[1] or '')
                if '{{MANDANT_ADRESSE}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{MANDANT_ADRESSE}}', mandant[2] or '')
                # Weitere Platzhalter hinzufügen...
        else:
            # Standard-Rechnung erstellen
            doc.add_paragraph(f"{firma[1]}")
            doc.add_paragraph(f"{firma[2]}")
            if firma[3]:
                doc.add_picture(firma[3], width=Inches(1.0))

            doc.add_paragraph(f"Mandant: {mandant[1]}")
            doc.add_paragraph(f"{mandant[2]}")

            doc.add_paragraph(f"Zeitraum: {rechnung[2]} bis {rechnung[3]}")

            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Beschreibung'
            hdr_cells[1].text = 'Anzahl'
            hdr_cells[2].text = 'Betrag'
            hdr_cells[3].text = 'Gesamt'

            for p in positionen:
                row_cells = table.add_row().cells
                row_cells[0].text = p[0]
                row_cells[1].text = str(p[1])
                row_cells[2].text = f"{p[2]:.2f} €"
                row_cells[3].text = f"{p[1] * p[2]:.2f} €"

            doc.add_paragraph(f"Netto: {rechnung[4]:.2f} €")
            if rechnung[5] > 0:
                doc.add_paragraph(f"MwSt ({firma[4]}%): {rechnung[5]:.2f} €")
            else:
                doc.add_paragraph(firma[5])
            doc.add_paragraph(f"Brutto: {rechnung[6]:.2f} €")

        filename = f"Rechnung_{rechnung_id}.docx"
        doc.save(filename)
        messagebox.showinfo("Exportiert", f"Rechnung als {filename} gespeichert")

    # Statistik
    def init_statistik_tab(self):
        ttk.Label(self.tab_statistik, text="Zeitraum:").grid(row=0, column=0)
        self.zeitraum_var = tk.StringVar()
        zeitraum_combo = ttk.Combobox(self.tab_statistik, textvariable=self.zeitraum_var, values=["Monat", "Quartal", "Jahr"])
        zeitraum_combo.grid(row=0, column=1)
        zeitraum_combo.current(0), "Umsatz pro Mandant"

        ttk.Label(self.tab_statistik, text="Wert:").grid(row=1, column=0)
        self.zeitraum_wert = ttk.Entry(self.tab_statistik)
        self.zeitraum_wert.grid(row=1, column=1)

        ttk.Button(self.tab_statistik, text="Anzeigen", command=self.show_statistik).grid(row=2, column=0, columnspan=2)

        self.statistik_canvas = None

    def show_statistik(self):
        zeitraum = self.zeitraum_var.get()
        wert = self.zeitraum_wert.get()
        if not wert:
            messagebox.showerror("Fehler", "Wert eingeben")
            return

        # Abfrage je nach Zeitraum
        if zeitraum == "Monat":
            query = "SELECT strftime('%Y-%m', erstellt_am) as period, SUM(summe_brutto) FROM rechnungen WHERE strftime('%Y-%m', erstellt_am) = ? GROUP BY period"
        elif zeitraum == "Quartal":
            query = "SELECT strftime('%Y', erstellt_am) || '-Q' || ((strftime('%m', erstellt_am) - 1) / 3 + 1) as period, SUM(summe_brutto) FROM rechnungen WHERE strftime('%Y', erstellt_am) = ? GROUP BY period"
        elif zeitraum == "Jahr":
            query = "SELECT strftime('%Y', erstellt_am) as period, SUM(summe_brutto) FROM rechnungen WHERE strftime('%Y', erstellt_am) = ? GROUP BY period"

        self.c.execute(query, (wert,))
        data = self.c.fetchall()

        if not data:
            messagebox.showinfo("Keine Daten", "Keine Rechnungen im Zeitraum")
            return

        # Plot
        if self.statistik_canvas:
            self.statistik_canvas.get_tk_widget().destroy()

        fig = plt.Figure(figsize=(5,4), dpi=100)
        ax = fig.add_subplot(111)
        periods = [row[0] for row in data]
        sums = [row[1] for row in data]
        ax.bar(periods, sums)
        ax.set_title(f"Umsatz {zeitraum} {wert}")
        ax.set_ylabel("Umsatz (€)")

        self.statistik_canvas = FigureCanvasTkAgg(fig, master=self.tab_statistik)
        self.statistik_canvas.get_tk_widget().grid(row=3, column=0, columnspan=2)

if __name__ == "__main__":
    root = tk.Tk()
    app = Rechnungsprogramm(root)
    root.mainloop()